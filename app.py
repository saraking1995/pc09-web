import io, os
from datetime import datetime
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook

DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
os.makedirs(DATA_DIR, exist_ok=True)

def find_data(name: str) -> str:
    path = os.path.join(DATA_DIR, name)
    if not os.path.exists(path):
        st.error(f"Không tìm thấy file trong /data: {name}")
    return path

def replace_text_placeholders(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        for r in p.runs:
            for k, v in mapping.items():
                if k in r.text:
                    r.text = r.text.replace(k, v)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        for k, v in mapping.items():
                            if k in r.text:
                                r.text = r.text.replace(k, v)

def insert_noidungbananh_from_excel(doc: Document, bookmark_name: str, excel_bytes: bytes) -> bool:
    wb = load_workbook(io.BytesIO(excel_bytes), data_only=True)
    ws = wb.active
    rows = []
    r = 3
    while True:
        a = ws[f"A{r}"].value or ""
        b = ws[f"B{r}"].value or ""
        c = ws[f"C{r}"].value or ""
        if not (a or b or c):
            break
        rows.append((str(a), str(b), str(c)))
        r += 1

    for para in doc.paragraphs:
        for bkm in para._element.xpath(".//w:bookmarkStart"):
            name = bkm.get(qn('w:name'))
            if name == bookmark_name:
                parent = para._element.getparent()
                idx = parent.index(para._element)

                def run_tx(t, bold=False):
                    rnode = OxmlElement('w:r')
                    tnode = OxmlElement('w:t'); tnode.text = t
                    rnode.append(tnode)
                    rPr = OxmlElement('w:rPr')
                    rf = OxmlElement('w:rFonts')
                    rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman')
                    rPr.append(rf)
                    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '28')
                    szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), '28')
                    rPr.append(sz); rPr.append(szcs)
                    if bold:
                        btag = OxmlElement('w:b'); rPr.append(btag)
                    rnode.insert(0, rPr)
                    return rnode

                for a, b, c in rows:
                    new_p = OxmlElement('w:p')
                    pPr = OxmlElement('w:pPr')
                    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'both')
                    pPr.append(jc); new_p.append(pPr)
                    new_p.append(run_tx(a, bold=True))
                    new_p.append(run_tx(b, bold=True))
                    rspace = OxmlElement('w:r')
                    tsp = OxmlElement('w:t'); tsp.set(qn('xml:space'), 'preserve'); tsp.text = " "
                    rspace.append(tsp); new_p.append(rspace)
                    new_p.append(run_tx(c, bold=False))
                    idx += 1
                    parent.insert(idx, new_p)
                return True
    return False

def build_bbgn_kemtheo(doc: Document, kemtheo_lines):
    target_idx = None
    parent = None
    for p in doc.paragraphs:
        if "kemtheo" in p.text.lower():
            parent = p._element.getparent()
            target_idx = parent.index(p._element)
            parent.remove(p._element)
            break
    if target_idx is None:
        return False

    def add_line(content: str | None, only_leader=False):
        new_para = doc.add_paragraph()
        new_para.paragraph_format.space_before = Pt(0)
        new_para.paragraph_format.space_after = Pt(0)
        new_para.paragraph_format.line_spacing = Pt(14.5)
        pPr = new_para._element.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right'); tab.set(qn('w:leader'), 'dot'); tab.set(qn('w:pos'), '9360')
        tabs.append(tab); pPr.append(tabs)

        if only_leader:
            r = new_para.add_run("\t"); r.font.name = 'Times New Roman'; r.font.size = Pt(8)
        else:
            r1 = new_para.add_run(content or "")
            r1.font.name = 'Times New Roman'; r1.font.size = Pt(13)
            r2 = new_para.add_run("\t")
            r2.font.name = 'Times New Roman'; r2.font.size = Pt(8)
        parent.insert(target_idx, new_para._element)

    if kemtheo_lines:
        kemtheo_lines = [line.strip() for line in kemtheo_lines if line.strip()]
        if kemtheo_lines:
            kemtheo_lines[-1] = kemtheo_lines[-1].rstrip(".") + "./."
    for line in kemtheo_lines:
        add_line(line, only_leader=False)
        target_idx += 1
    add_line(None, only_leader=True)
    return True

st.set_page_config(page_title="Hỗ Trợ Làm Bản Ảnh - Web", layout="centered")
st.title("Hỗ Trợ Làm Bản Ảnh • Bản web (Streamlit)")

with st.expander("📘 Tải file mẫu TMBA.xlsm"):
    tmba_path = find_data("TMBA.xlsm")
    if os.path.exists(tmba_path):
        with open(tmba_path, "rb") as f:
            st.download_button("Tải TMBA.xlsm", data=f.read(), file_name="TMBA.xlsm",
                               mime="application/vnd.ms-excel.sheet.macroEnabled.12")

tabs = st.tabs(["📄 Tạo THUYẾT MINH", "📑 Tạo BBGN"])

with tabs[0]:
    st.subheader("Tạo Thuyết minh")
    col1, col2 = st.columns(2)
    loaibananh = col1.selectbox("Loại bản ảnh", [
        "khám nghiệm hiện trường",
        "khám nghiệm tử thi",
        "khám nghiệm hiện trường và tử thi",
        "khám phương tiện liên quan đến tai nạn giao thông",
        "Tùy chỉnh"
    ])
    loaibananh_custom = ""
    if loaibananh == "Tùy chỉnh":
        loaibananh_custom = col2.text_input("Nhập loại bản ảnh tùy chỉnh", "")
    vu_viec = st.selectbox("Tên vụ việc", [
        "Chết người","Giết người","Cố ý gây thương tích","Tai nạn giao thông đường bộ",
        "Hủy hoại tài sản","Trộm cắp tài sản","Hiếp dâm","Tùy chỉnh"
    ])
    vu_viec_custom = ""
    if vu_viec == "Tùy chỉnh":
        vu_viec_custom = st.text_input("Nhập tên vụ việc tùy chỉnh", "")

    xrph = st.radio("Sự kiện", ["xảy ra","phát hiện"], horizontal=True)
    nkn = st.date_input("Ngày KN", format="DD/MM/YYYY")
    nxr = st.date_input("Ngày xảy ra", format="DD/MM/YYYY")
    diadiem = st.text_input("Địa điểm xảy ra", "")
    uploaded_tmba = st.file_uploader("Tải lên TMBA.xlsm", type=["xlsm"])

    if st.button("Sinh Thuyết minh (DOCX)", use_container_width=True, type="primary"):
        if uploaded_tmba is None:
            st.error("Vui lòng tải lên TMBA.xlsm.")
        else:
            doc = Document(find_data("Mau_TM.docx"))
            mapping = {
                "loaibananh": (loaibananh_custom if loaibananh == "Tùy chỉnh" else loaibananh),
                "nkn": nkn.strftime("%d/%m/%Y"),
                "vuviec": (vu_viec_custom if vu_viec == "Tùy chỉnh" else vu_viec),
                "xrph": xrph,
                "nxr": nxr.strftime("%d/%m/%Y"),
                "dd": diadiem
            }
            replace_text_placeholders(doc, mapping)
            _ = insert_noidungbananh_from_excel(doc, "noidungbananh", uploaded_tmba.read())
            buf = io.BytesIO(); doc.save(buf); buf.seek(0)
            outname = f"TMBA_{(mapping['vuviec'] or 'VU_VIEC').replace(' ','_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            st.success("Đã tạo thuyết minh.")
            st.download_button("Tải file Thuyết minh", data=buf.getvalue(), file_name=outname,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tabs[1]:
    st.subheader("Tạo Biên bản giao nhận (BBGN)")
    vuviec_b = st.text_input("Tên vụ việc", "")
    xrph_b = st.radio("Sự kiện", ["xảy ra","phát hiện"], horizontal=True, index=0)
    diadiem_b = st.text_input("Địa điểm xảy ra", "")
    nxr_b = st.date_input("Ngày xảy ra", format="DD/MM/YYYY", key="nxr_b")
    kemtheo_text = st.text_area("Danh mục kèm theo (mỗi dòng 1 mục)",
                                "- 01 (một) Biên bản khám nghiệm hiện trường.")
    if st.button("Sinh BBGN (DOCX)", use_container_width=True):
        doc = Document(find_data("BBGN.docx"))
        replace_text_placeholders(doc, {
            "vuviec": vuviec_b, "xrph": xrph_b,
            "diadiem": diadiem_b, "nxr": nxr_b.strftime("%d/%m/%Y")
        })
        kem_lines = kemtheo_text.splitlines()
        _ = build_bbgn_kemtheo(doc, kem_lines)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        fname = f"BienBanGiaoNhan_{(vuviec_b or 'VU_VIEC').replace(' ','_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        st.success("Đã tạo BBGN.")
        st.download_button("Tải file BBGN", data=buf.getvalue(), file_name=fname,
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.caption("© PC09 Khánh Hòa")
