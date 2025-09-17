import os
import shutil
import re
from zipfile import ZipFile
from openpyxl import load_workbook
from lxml import etree
from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import logging
import sys
from PIL import Image

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)  # dùng stdout cho hệ thống log GUI bắt được
    ]
)
log = logging.getLogger(__name__)

# ==== CẤU HÌNH ==== 

template_dir = r"D:\PC09\data\Mau PC09"
temp_output_dir = r"D:\PC09\TempTach"
final_output_dir = r"D:\PC09\KetQuaWord"
os.makedirs(temp_output_dir, exist_ok=True)
os.makedirs(final_output_dir, exist_ok=True)

# Chọn thư mục chứa TMBA.xlsm
folder = os.environ.get("EXPORT_FOLDER")
if not folder or not os.path.exists(os.path.join(folder, "TMBA.xlsm")):
    log.warning(" Không tìm thấy thư mục chứa TMBA.xlsm được truyền từ GUI.")
    exit()

image_dir = os.environ.get("EXPORT_FOLDER")
if not image_dir or not os.path.isdir(image_dir):
    log.warning("❌ Không tìm thấy thư mục ảnh hợp lệ.")
    exit()

# ==== ĐỌC EXCEL ==== 
excel_path = os.path.join(folder, "TMBA.xlsm")
if not os.path.exists(excel_path):
    log.warning(" Không tìm thấy TMBA.xlsm trong thư mục đã chọn.")
    exit()


wb = load_workbook(excel_path)
ws = wb.active

# ==== LẤY SỐ ẢNH LỚN NHẤT TỪ CỘT A ====
so_anh_toan_bo = 0
for row in range(3, ws.max_row + 1):
    ten_anh = ws[f"A{row}"].value
    if not ten_anh:
        break
    nums = re.findall(r"\d+", str(ten_anh))
    if nums:
        idx = int(nums[0])
        if idx > so_anh_toan_bo:
            so_anh_toan_bo = idx

# ==== LẤY TÊN FILE TỔNG HỢP ==== 
ten_file_output = ws["D1"].value or "TongHop"
ten_file_output = str(ten_file_output).replace("/", "-").replace("\\", "-").strip()
final_output = os.path.join(final_output_dir, ten_file_output + ".docx")


def extract_image_number(text):
    nums = re.findall(r"\d+", text)
    return int(nums[0]) if nums else None

def create_HH_doc(template_path, save_path, so_anh_toan_bo):
    temp_dir = os.path.join(temp_output_dir, "temp_HH")
    os.makedirs(temp_dir, exist_ok=True)

    with ZipFile(template_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    doc_path = os.path.join(temp_dir, 'word', 'document.xml')
    tree = etree.parse(doc_path)
    root = tree.getroot()
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    ngay_hoan_thanh = datetime.today().strftime("%d/%m/%Y")

    # Tìm và thay nội dung
    for p in root.findall(".//w:p", namespaces=ns):
        for t in p.findall(".//w:t", namespaces=ns):
            if t.text:
                if "BẢN ẢNH GỒM" in t.text:
                    t.text = f"BẢN ẢNH GỒM: {so_anh_toan_bo:02} ẢNH CỠ (10X15)CM"
                elif "HOÀN THÀNH NGÀY" in t.text:
                    t.text = f"HOÀN THÀNH NGÀY: {ngay_hoan_thanh}"

    tree.write(doc_path, xml_declaration=True, encoding='utf-8', standalone='yes')

    zip_base = os.path.join(temp_output_dir, "temp_HH_output")
    archive_path = shutil.make_archive(zip_base, 'zip', temp_dir)
    shutil.move(archive_path, save_path)
    shutil.rmtree(temp_dir)


def create_word_from_template(template_path, save_path, texts, image_indices, image_dir, so_anh_toan_bo):
    temp_dir = os.path.join(temp_output_dir, "temp_extract")
    os.makedirs(temp_dir, exist_ok=True)
    template_base = os.path.basename(template_path).upper()
    ngay_hoan_thanh = datetime.today().strftime("%d/%m/%Y")  
    
    with ZipFile(template_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    # === GHI TEXTBOX === 
    doc_path = os.path.join(temp_dir, 'word', 'document.xml')
    tree = etree.parse(doc_path)
    root = tree.getroot()
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    tb_contents = root.findall('.//w:txbxContent', namespaces=ns)

    for target_index, tb in enumerate(tb_contents):

        if target_index < len(texts):
            for p in tb.findall('.//w:p', namespaces=ns):
                p.getparent().remove(p)

            empty_p = OxmlElement('w:p')
            tb.append(empty_p)

            new_p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'both')
            pPr.append(jc)
            new_p.append(pPr)

            ten_anh_raw, sep, mo_ta = texts[target_index].partition(":")
            ten_anh = ten_anh_raw.strip() + ":"

            r_bold = OxmlElement('w:r')
            rPr_bold = OxmlElement('w:rPr')
            rPr_bold.append(OxmlElement('w:b'))
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '28')
            rPr_bold.append(sz)
            rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:ascii'), 'Times New Roman')
            rPr_bold.append(rFonts)
            t_bold = OxmlElement('w:t'); t_bold.set(qn('xml:space'), 'preserve'); t_bold.text = ten_anh + " "
            r_bold.append(rPr_bold); r_bold.append(t_bold)

            r_normal = OxmlElement('w:r')
            t_normal = OxmlElement('w:t'); t_normal.text = mo_ta.strip()
            r_normal.append(t_normal)
            rPr_normal = OxmlElement('w:rPr')
            sz_n = OxmlElement('w:sz'); sz_n.set(qn('w:val'), '28')
            rPr_normal.append(sz_n)
            rFonts_n = OxmlElement('w:rFonts'); rFonts_n.set(qn('w:ascii'), 'Times New Roman')
            rPr_normal.append(rFonts_n)
            r_normal.insert(0, rPr_normal)

            new_p.append(r_bold); new_p.append(r_normal)
            tb.append(new_p)
    
    template_base = os.path.basename(template_path).upper()
    if "N-H" in template_base or "D-H" in template_base:
    # Thay nội dung BẢN ẢNH GỒM và NGÀY HOÀN THÀNH
        for p in root.findall(".//w:p", namespaces=ns):
            texts_in_p = p.findall(".//w:t", namespaces=ns)
            for t in texts_in_p:
                if t.text and "BẢN ẢNH GỒM" in t.text:
                    t.text = f"BẢN ẢNH GỒM: {so_anh_toan_bo:02} ẢNH CỠ (10X15)CM"
                if t.text and "HOÀN THÀNH NGÀY" in t.text:
                    t.text = f"HOÀN THÀNH NGÀY: {ngay_hoan_thanh}"

    tree.write(doc_path, xml_declaration=True, encoding='utf-8', standalone='yes')

    # === THAY ẢNH === 
    media_path = os.path.join(temp_dir, 'word', 'media')
    template_base = os.path.basename(template_path).upper()
    ngay_hoan_thanh = datetime.now().strftime("%d/%m/%Y")

    image_paths = [os.path.join(image_dir, f"{idx}.jpg") for idx in image_indices if idx]

    def replace_images_for_DN_DD(media_path, image_paths):
        from PIL import Image
        for i, img_path in enumerate(image_paths[:2]):
            if not os.path.exists(img_path):
                continue
            with Image.open(img_path) as img:
                width, height = img.size
                if height > width:  # ảnh đứng
                    img = img.rotate(270, expand=True)
                img.save(os.path.join(media_path, f"image{i+1}.jpeg"), format='JPEG', quality=95)

    if "D-N" in template_base or "D-D" in template_base:
        replace_images_for_DN_DD(media_path, image_paths)
    elif "N-H" in template_base or "D-H" in template_base:
        if len(image_paths) >= 1:
            shutil.copy(image_paths[0], os.path.join(media_path, "image1.jpeg"))
    else:
        if len(image_paths) >= 2:
            shutil.copy(image_paths[0], os.path.join(media_path, "image1.jpeg"))
            shutil.copy(image_paths[1], os.path.join(media_path, "image2.jpeg"))


    # === LƯU FILE === 
    zip_base = os.path.join(temp_output_dir, "temp_output")
    archive_path = shutil.make_archive(zip_base, 'zip', temp_dir)
    shutil.move(archive_path, save_path)
    shutil.rmtree(temp_dir)


# ==== XỬ LÝ EXCEL VÀ TẠO FILE ==== 
row = 3
temp_files = []

while True:
    ten_anh_1 = ws[f"A{row}"].value
    mo_ta_1 = ws[f"B{row}"].value or ""
    mo_ta_2 = ws[f"C{row}"].value or ""
    ma_mau = ws[f"D{row}"].value

    if not ten_anh_1 or not ma_mau:
        break

    ten_anh_2 = ws[f"A{row+1}"].value
    mo_ta_1_b = ws[f"B{row+1}"].value or ""
    mo_ta_2_b = ws[f"C{row+1}"].value or ""

    text1 = f"{ten_anh_1} {mo_ta_1} {mo_ta_2}"
    text2 = f"{ten_anh_2} {mo_ta_1_b} {mo_ta_2_b}" if ten_anh_2 else ""

    if ten_anh_2:
        # Có 2 ảnh
        if ma_mau in ["N-H", "D-D", "D-H", "N-D"]:
            texts = [text1, text2]
        elif ma_mau == "N-N":
            texts = [text2, text1]  # Ảnh 03 → mô tả 03, Ảnh 04 → mô tả 04 (không đảo)
        elif ma_mau in ["D-N"]:
            texts = [text2, text1]
        else:
            texts = [text1, text2]
        if ma_mau == "N-N":
            image_indices = [
                extract_image_number(ten_anh_2),
                extract_image_number(ten_anh_1)
            ]
        elif ma_mau in ["N-D", "D-D"]:
            image_indices = [
                extract_image_number(ten_anh_2),
                extract_image_number(ten_anh_1)   
            ]
        else:
            image_indices = [
                extract_image_number(ten_anh_1),
                extract_image_number(ten_anh_2)
            ]
        row_step = 2
    else:
        
        texts = [text1]
        image_indices = [extract_image_number(ten_anh_1)]
        row_step = 1

    template_file = os.path.join(template_dir, f"{ma_mau}.docx")
    temp_doc_path = os.path.join(temp_output_dir, f"temp_{row}_{ma_mau}.docx")

    if not os.path.exists(template_file):
        log.warning(f"Không tìm thấy mẫu: {template_file}")
        row += row_step
        continue

    create_word_from_template(template_file, temp_doc_path, texts, image_indices, image_dir, so_anh_toan_bo)

    try:
        so_thutu = int(re.findall(r"\d+", ten_anh_1)[0])
    except:
        so_thutu = row

    if os.path.exists(temp_doc_path):
        temp_files.append((so_thutu, temp_doc_path))

    row += row_step

# ==== KIỂM TRA MẪU CUỐI CÙNG ====
last_ma_mau = None
for row_check in range(3, ws.max_row + 1):
    ma_mau_check = ws[f"D{row_check}"].value
    if ma_mau_check:
        last_ma_mau = ma_mau_check.strip().upper()

them_HH_sau_cung = (last_ma_mau == "H-H")

# ==== KIỂM TRA MẪU CUỐI CÓ PHẢI H-H KHÔNG ==== 
them_HH_sau_cung = False
for row_back in range(ws.max_row, 2, -1):  # từ dưới lên trên
    val = ws[f"D{row_back}"].value
    if val and str(val).strip():
        if str(val).strip().upper() == "H-H":
            them_HH_sau_cung = True
        break  # chỉ xét ô cuối cùng có dữ liệu ở cột D

loai_val = os.environ.get("Loai_ban_anh", "")
vuviec_val = os.environ.get("vuviec", "")
xrph_val = os.environ.get("xrph", "")
diadiem_val = os.environ.get("diadiem", "")
nxr_val = os.environ.get("nxr", "")
nkn_val = os.environ.get("nkn", "")
ngaykn_val = ""
thangkn_val = ""
try:
    d = datetime.strptime(nkn_val, "%d/%m/%Y")
    ngaykn_val = str(d.day).zfill(2)
    thangkn_val = str(d.month).zfill(2)
except:
    log.warning("⚠️ Không thể phân tích ngày từ nkn_val: " + str(nkn_val))
ngay_hoan_thanh = datetime.now().strftime("%d/%m/%Y")
ban_anh_path = os.path.join(temp_output_dir, "Ban_anh.docx")

def generate_ban_anh(loai_val, vuviec_val, xrph_val, diadiem_val, nxr_val, ngaykn_val, thangkn_val, ngay_lap_val, so_anh_val, save_path):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    template_path = r"D:/PC09/data/BanAnhTemplate.docx"
    doc = Document(template_path)
    # Thiết lập lề trái 2.6 cm
    section = doc.sections[0]
    section.left_margin = Cm(2.6)

    def replace_in_doc(doc_or_cell, find_text, replace_text, tableader_fontsize=None):
        for p in doc_or_cell.paragraphs:
            if find_text in p.text:
                p.paragraph_format.line_spacing = Pt(18)  # dãn dòng 1.15 cho font 14pt

                for i, run in enumerate(p.runs):
                    if find_text in run.text:
                        parts = run.text.split(find_text)
                        before = parts[0]
                        after = parts[1] if len(parts) > 1 else ""

                        run.text = before + replace_text
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        if "\t" in after:
                            tab_run = p.add_run("\t")
                            tab_run.font.name = 'Times New Roman'
                            tab_run.font.size = Pt(8)
                            remain = after.replace("\t", "")
                            if remain:
                                tail_run = p.add_run(remain)
                                tail_run.font.name = 'Times New Roman'
                                tail_run.font.size = Pt(14)
                        else:
                            run.text += after
                break
        # Đệ quy vào bảng nếu có
        if hasattr(doc_or_cell, "tables"):
            for table in doc_or_cell.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_doc(cell, find_text, replace_text, tableader_fontsize)

    # === Gọi thay thế ===
    replace_in_doc(doc, "loai_ban_anh", loai_val)
    replace_in_doc(doc, "vuviec", vuviec_val)
    replace_in_doc(doc, "xrph", xrph_val)
    replace_in_doc(doc, "diadiem", diadiem_val)
    replace_in_doc(doc, "nxr", nxr_val)
    replace_in_doc(doc, "ngaykn", ngaykn_val)
    replace_in_doc(doc, "thangkn", thangkn_val)
    replace_in_doc(doc, "ngay_hoan_thanh", ngay_lap_val)
    replace_in_doc(doc, "so_anh", so_anh_val)

    doc.save(save_path)
    log.info(f"📄 Đã tạo Ban_anh.docx tại: {save_path}")


generate_ban_anh(
    loai_val, vuviec_val, xrph_val, diadiem_val,
    nxr_val, ngaykn_val, thangkn_val, ngay_hoan_thanh,
    f"{so_anh_toan_bo:02}", ban_anh_path
)

# ==== GỘP FILE ==== 
if temp_files:
    temp_files.sort(key=lambda x: x[0])
    master = Document(temp_files[0][1])
    composer = Composer(master)

    for _, temp_path in temp_files[1:]:
        blank = Document(); blank.add_page_break()
        composer.append(blank)
        composer.append(Document(temp_path))

    # === CHÈN H-H VÀO CUỐI NẾU CÓ ===
    if them_HH_sau_cung:
        hh_template = os.path.join(template_dir, "H-H.docx")
        hh_temp_doc = os.path.join(temp_output_dir, "temp_HH.docx")
        if os.path.exists(hh_template):
            create_HH_doc(hh_template, hh_temp_doc, so_anh_toan_bo)

            # ➕ Thêm trang trắng trước khi chèn H-H
            blank = Document()
            blank.add_page_break()
            composer.append(blank)
            composer.append(Document(hh_temp_doc))
            log.info("📎 Đã chèn mẫu H-H với nội dung cập nhật vào cuối văn bản.")
        else:
            log.warning("⚠️ Không tìm thấy file H-H.docx để chèn vào cuối.")
    
        # ➕ Chèn thêm Ban_anh.docx vào cuối nếu có
    if os.path.exists(ban_anh_path):
        blank = Document()
        blank.add_page_break()
        composer.append(blank)
        composer.append(Document(ban_anh_path))
        log.info("📎 Đã chèn Ban_anh.docx vào cuối văn bản.")
    else:
        log.warning("⚠️ Không tìm thấy Ban_anh.docx để chèn vào cuối.")


    # ✅ LUÔN LUÔN LƯU FILE Ở ĐÂY — KHÔNG ĐẶT TRONG IF
    composer.save(final_output)
    shutil.rmtree(temp_output_dir)
    log.warning(f"\n✅ Đã tạo file tổng hợp tại: {final_output}")

else:
    log.warning(f"❌ Không có file nào được tạo.")

if __name__ == "__main__":
    pass
    # toàn bộ xử lý chính nằm ở đây