from fastapi import FastAPI, UploadFile, Form
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
import io, os
from datetime import datetime

app = FastAPI()

DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data")

def replace_text(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        for r in p.runs:
            for k, v in mapping.items():
                if k in r.text:
                    r.text = r.text.replace(k, v)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        for k, v in mapping.items():
                            if k in r.text:
                                r.text = r.text.replace(k, v)

def insert_from_excel(doc: Document, bookmark_name: str, file_bytes: bytes):
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    rows = []
    r = 3
    while True:
        a = ws[f"A{r}"].value or ""
        b = ws[f"B{r}"].value or ""
        c = ws[f"C{r}"].value or ""
        if not (a or b or c):
            break
        rows.append((str(a), str(b), str(c)))
        r += 1
    for para in doc.paragraphs:
        for bkm in para._element.xpath(".//w:bookmarkStart"):
            if bkm.get(qn('w:name')) == bookmark_name:
                parent = para._element.getparent()
                idx = parent.index(para._element)
                for a, b, c in rows:
                    new_p = OxmlElement('w:p')
                    pPr = OxmlElement('w:pPr')
                    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'both')
                    pPr.append(jc); new_p.append(pPr)

                    def run_tx(t, bold=False):
                        rnode = OxmlElement('w:r')
                        tnode = OxmlElement('w:t'); tnode.text = t
                        rnode.append(tnode)
                        rPr = OxmlElement('w:rPr')
                        rf = OxmlElement('w:rFonts')
                        rf.set(qn('w:ascii'), 'Times New Roman')
                        rf.set(qn('w:hAnsi'), 'Times New Roman')
                        rPr.append(rf)
                        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '28')
                        szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), '28')
                        rPr.append(sz); rPr.append(szcs)
                        if bold: rPr.append(OxmlElement('w:b'))
                        rnode.insert(0, rPr)
                        return rnode

                    new_p.append(run_tx(a, bold=True))
                    new_p.append(run_tx(b, bold=True))
                    rspace = OxmlElement('w:r')
                    tsp = OxmlElement('w:t'); tsp.set(qn('xml:space'), 'preserve'); tsp.text = " "
                    rspace.append(tsp); new_p.append(rspace)
                    new_p.append(run_tx(c, bold=False))
                    idx += 1
                    parent.insert(idx, new_p)
                return True
    return False

@app.post("/thuyetminh")
async def tao_thuyet_minh(
    loaibananh: str = Form(...),
    vuviec: str = Form(...),
    xrph: str = Form(...),
    nkn: str = Form(...),
    nxr: str = Form(...),
    diadiem: str = Form(...),
    tmba: UploadFile = None
):
    template = os.path.join(DATA_DIR, "Mau_TM.docx")
    doc = Document(template)
    mapping = {
        "loaibananh": loaibananh,
        "vuviec": vuviec,
        "xrph": xrph,
        "nkn": nkn,
        "nxr": nxr,
        "dd": diadiem
    }
    replace_text(doc, mapping)
    if tmba:
        insert_from_excel(doc, "noidungbananh", await tmba.read())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    outname = f"TMBA_{vuviec}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": f"attachment; filename={outname}"})

@app.post("/bbgn")
async def tao_bbgn(
    vuviec: str = Form(...),
    xrph: str = Form(...),
    nxr: str = Form(...),
    diadiem: str = Form(...),
    kemtheo: str = Form("Không")
):
    template = os.path.join(DATA_DIR, "BBGN.docx")
    doc = Document(template)
    replace_text(doc, {
        "vuviec": vuviec,
        "xrph": xrph,
        "nxr": nxr,
        "diadiem": diadiem
    })
    # Đơn giản hóa: chỉ chèn text kemtheo vào cuối tài liệu
    doc.add_paragraph("Kèm theo:")
    for line in kemtheo.splitlines():
        doc.add_paragraph(f"- {line}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    outname = f"BBGN_{vuviec}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": f"attachment; filename={outname}"})

